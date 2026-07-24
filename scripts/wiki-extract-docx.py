#!/usr/bin/env python3
# /// script
# requires-python = ">=3.10"
# dependencies = ["python-docx>=1.1"]
# ///
"""
wiki-extract-docx.py — extract markdown from a Word document via python-docx.

Stage 12 of the wiki upgrade arc — see [[Eric Ma — Mastering PKM with Obsidian
and AI]] for the canonical recipe.

Usage:
  uv run wiki-extract-docx.py <path-to-docx>      # markdown to stdout
  uv run wiki-extract-docx.py --self-test          # synthesize tmp DOCX, parse, exit 0/1

LIMITATIONS:
- Tracked changes, comments, embedded objects not extracted.
- Custom styles map to nearest standard markdown analog (Heading 1/2/3 → ##/###/####).
- Complex tables (merged cells, nested) may degrade.

First run cost: ~30-60 sec (uv resolves python-docx + lxml wheels). Cached after.
"""
import sys
import tempfile
from pathlib import Path

# Force UTF-8 on stdout/stderr so non-ASCII content (accented chars, em-dashes,
# math symbols) doesn't crash on Windows where the default codec is cp1252.
# Python 3.7+ feature; harmless on Linux/macOS.
try:
    sys.stdout.reconfigure(encoding="utf-8", errors="replace")
    sys.stderr.reconfigure(encoding="utf-8", errors="replace")
except (AttributeError, ValueError):
    pass

try:
    from docx import Document
    from docx.oxml.ns import qn
    from docx.text.paragraph import Paragraph
    from docx.table import Table
except ImportError:
    sys.stderr.write(
        "missing dep: python-docx. Run via `uv run wiki-extract-docx.py ...` "
        "(PEP-723 inline deps), not bare python.\n"
    )
    sys.exit(1)


def _paragraph_md(para):
    text = para.text.strip()
    if not text:
        return ""
    style = (para.style.name or "").lower() if para.style else ""
    if "heading 1" in style:
        return f"## {text}"
    if "heading 2" in style:
        return f"### {text}"
    if "heading 3" in style:
        return f"#### {text}"
    if "heading 4" in style:
        return f"##### {text}"
    if "title" in style:
        return f"# {text}"
    if "list" in style or "bullet" in style:
        return f"- {text}"
    return text


def _table_md(tbl):
    rows = [[cell.text.strip().replace("\n", " ").replace("|", "\\|") for cell in row.cells] for row in tbl.rows]
    if not rows:
        return ""
    cols = max(len(r) for r in rows)
    rows = [r + [""] * (cols - len(r)) for r in rows]
    head = "| " + " | ".join(rows[0]) + " |"
    sep = "| " + " | ".join(["---"] * cols) + " |"
    body = ["| " + " | ".join(r) + " |" for r in rows[1:]]
    return "\n".join([head, sep] + body)


def doc_to_markdown(doc):
    """Walk doc body in document order, dispatching paragraphs and tables."""
    out = []
    body = doc.element.body
    for child in body.iterchildren():
        if child.tag == qn("w:p"):
            md = _paragraph_md(Paragraph(child, doc))
            if md:
                out.append(md)
        elif child.tag == qn("w:tbl"):
            md = _table_md(Table(child, doc))
            if md:
                out.append(md)
    return "\n\n".join(out)


def self_test():
    KNOWN = "Stage 12 docx selftest marker"
    KNOWN_HEAD = "Self Test Heading"
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
        tmp_path = Path(f.name)
    try:
        doc = Document()
        doc.add_heading(KNOWN_HEAD, level=1)
        doc.add_paragraph(KNOWN)
        tbl = doc.add_table(rows=2, cols=2)
        tbl.rows[0].cells[0].text = "h1"
        tbl.rows[0].cells[1].text = "h2"
        tbl.rows[1].cells[0].text = "v1"
        tbl.rows[1].cells[1].text = "v2"
        doc.save(str(tmp_path))

        rt = Document(str(tmp_path))
        md = doc_to_markdown(rt)

        if KNOWN not in md:
            sys.stderr.write(f"selftest FAIL: marker not in extracted output\n  got:\n{md}\n")
            sys.exit(1)
        if f"## {KNOWN_HEAD}" not in md:
            sys.stderr.write(f"selftest FAIL: heading not detected\n  got:\n{md}\n")
            sys.exit(1)
        if "| h1 | h2 |" not in md:
            sys.stderr.write(f"selftest FAIL: table not extracted\n  got:\n{md}\n")
            sys.exit(1)
        sys.stdout.write("selftest pass\n")
        sys.exit(0)
    finally:
        try:
            tmp_path.unlink()
        except OSError:
            pass


def main():
    if len(sys.argv) >= 2 and sys.argv[1] == "--self-test":
        return self_test()
    if len(sys.argv) < 2:
        sys.stderr.write("usage: uv run wiki-extract-docx.py <path-to-docx> | --self-test\n")
        sys.exit(2)

    path = Path(sys.argv[1])
    if not path.exists():
        sys.stderr.write(f"file not found: {path}\n")
        sys.exit(1)

    try:
        doc = Document(str(path))
    except Exception as e:
        sys.stderr.write(f"python-docx failed: {type(e).__name__}: {e}\n")
        sys.exit(1)

    out = [f"# {path.stem}", "", f"_Source: `{path.name}`_", "", doc_to_markdown(doc), ""]
    sys.stdout.write("\n".join(out))


if __name__ == "__main__":
    main()
